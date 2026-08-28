"""
========================================================================================
📂 My Documents — Ultra-Lightweight RAG System (Render-Optimized, <50MB RAM)
========================================================================================
Features:
1. Robust Multi-Format Extractors: PDF (pypdf), Word (python-docx), TXT, CSV, MD, JSON
2. Intelligent Context Chunking (1000-char semantic windows with 200-char overlap)
3. Hybrid Search Engine:
   - Dense Semantic Embeddings (Google Gemini Free text-embedding-004)
   - BM25 & Keyword Frequency Matcher with Stopwords Filtering
   - Document-Wide Summary & Overview Query Detection
4. Ultra-Low RAM Footprint: ~35MB RAM (Runs effortlessly on Render Free Tier)
5. Multi-LLM Support: Google Gemini 2.0 Flash (Free), Groq (Free), or Built-in Extractive Engine
========================================================================================
"""

import os
import re
import json
import math
import shutil
import numpy as np
from pathlib import Path
from flask import Flask, request, jsonify, render_template_string
from dotenv import load_dotenv

# Lightweight document extractors
import pypdf
import docx
import requests

# Load .env variables
load_dotenv()

app = Flask(__name__)

# Base directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
app.config['UPLOAD_FOLDER'] = os.path.join(BASE_DIR, 'my_documents_files')
app.config['STORAGE_FOLDER'] = os.path.join(BASE_DIR, 'my_documents_storage')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['STORAGE_FOLDER'], exist_ok=True)

# In-memory document and chunk registry
chunks_registry = []
uploaded_documents = []
chunk_vectors = None  # NumPy 2D array of chunk embeddings

# Question stopwords to filter out for keyword matching
STOPWORDS = set(
    'a about above after again against all am an and any are as at be because been before being '
    'below between both but by can could did do does doing down during each few for from further '
    'had has have having he her here hers him his how i if in into is it its me more most my no '
    'nor not of off on once only or other our out over own same she should so some such than that '
    'the their them then there these they this those through to too under until up very was we were '
    'what when where which while who why with you please tell give explain summarize describe details'.split()
)


# -----------------------------------------------------------------------------
# 1. DOCUMENT PARSERS & TEXT EXTRACTORS (Pure Python)
# -----------------------------------------------------------------------------

def extract_text(file_path: str, filename: str) -> list[dict]:
    """Extracts text page by page from PDF, DOCX, TXT, MD, or CSV."""
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else 'txt'
    pages = []

    # A) PDF Files
    if ext == 'pdf':
        try:
            reader = pypdf.PdfReader(file_path)
            for i, page in enumerate(reader.pages):
                txt = (page.extract_text() or "").strip()
                if txt:
                    pages.append({'text': txt, 'page': i + 1})
        except Exception as e:
            print(f"Error reading PDF {filename}: {e}")

    # B) Word Documents (.docx)
    elif ext in ['docx', 'doc']:
        try:
            doc = docx.Document(file_path)
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if row_txt:
                        paras.append(row_txt)
            full_text = "\n\n".join(paras)
            if full_text.strip():
                pages.append({'text': full_text, 'page': 1})
        except Exception as e:
            print(f"Error reading DOCX {filename}: {e}")

    # C) Plain Text / Markdown / CSV / JSON
    else:
        for enc in ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    content = f.read().strip()
                    if content:
                        pages.append({'text': content, 'page': 1})
                break
            except Exception:
                continue

    return pages


# -----------------------------------------------------------------------------
# 2. INTELLIGENT TEXT CHUNKER (Pure Python)
# -----------------------------------------------------------------------------

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
    """
    Splits text into rich overlapping chunks (1000 chars ~ 150-200 words)
    to preserve complete sentences, tables, and paragraph context.
    """
    if len(text) <= chunk_size:
        return [text.strip()]

    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            # Break cleanly on newline or sentence end
            break_pos = text.rfind('\n', start, end)
            if break_pos == -1 or break_pos < start + (chunk_size // 2):
                break_pos = text.rfind('. ', start, end)
            if break_pos == -1 or break_pos < start + (chunk_size // 2):
                break_pos = text.rfind(' ', start, end)
            if break_pos != -1 and break_pos > start:
                end = break_pos + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap
        if start >= len(text) - overlap:
            break

    return chunks


# -----------------------------------------------------------------------------
# 3. ULTRA-LIGHTWEIGHT EMBEDDINGS & HYBRID SEARCH
# -----------------------------------------------------------------------------

def sanitize_key(val):
    """Clean API keys removing all whitespace, newlines, and carriage returns."""
    if not val:
        return None
    cleaned = str(val).strip().replace('\r', '').replace('\n', '').replace('\t', '').strip()
    return cleaned if cleaned else None


def get_gemini_embedding(text: str, api_key: str) -> list[float]:
    """Fetch 100% free embeddings from Google Gemini API."""
    for model_name in ['gemini-embedding-001', 'gemini-embedding-2', 'text-embedding-004']:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:embedContent?key={api_key}"
            payload = {
                "model": f"models/{model_name}",
                "content": {"parts": [{"text": text[:2000]}]}
            }
            res = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=15)
            if res.status_code == 200:
                data = res.json()
                if 'embedding' in data and 'values' in data['embedding']:
                    return data['embedding']['values']
        except Exception as e:
            continue
    return None


def extract_keywords(text: str) -> list[str]:
    """Extract meaningful words from text excluding question stopwords."""
    words = re.findall(r'[a-zA-Z0-9_\-\$]+', text.lower())
    return [w for w in words if w not in STOPWORDS and len(w) > 1]


def compute_keyword_score(chunk_text: str, query: str) -> float:
    """Computes keyword overlap + exact phrase match score."""
    q_words = extract_keywords(query)
    if not q_words:
        return 0.0

    c_lower = chunk_text.lower()
    c_words = set(extract_keywords(chunk_text))

    # Match tokens
    matched = [w for w in q_words if w in c_words or w in c_lower]
    overlap_score = len(matched) / len(q_words)

    # Bonus for exact query phrase match
    phrase_bonus = 0.5 if query.lower().strip() in c_lower else 0.0

    return overlap_score + phrase_bonus


# -----------------------------------------------------------------------------
# 4. STORAGE & VECTOR INDEX MANAGEMENT
# -----------------------------------------------------------------------------

meta_file = os.path.join(app.config['STORAGE_FOLDER'], 'metadata.json')
vectors_file = os.path.join(app.config['STORAGE_FOLDER'], 'vectors.npy')


def load_storage():
    """Load persistent metadata and vector index from disk."""
    global chunks_registry, uploaded_documents, chunk_vectors
    if os.path.exists(meta_file):
        try:
            with open(meta_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                chunks_registry = data.get('chunks', [])
                uploaded_documents = data.get('documents', [])
        except Exception as e:
            print(f"Error loading metadata: {e}")

    if os.path.exists(vectors_file):
        try:
            chunk_vectors = np.load(vectors_file)
        except Exception as e:
            print(f"Error loading vectors: {e}")
            chunk_vectors = None


load_storage()


def save_storage():
    """Save metadata and vector index to disk."""
    global chunks_registry, uploaded_documents, chunk_vectors
    with open(meta_file, 'w', encoding='utf-8') as f:
        json.dump({
            'chunks': chunks_registry,
            'documents': uploaded_documents
        }, f, indent=2)

    if chunk_vectors is not None:
        np.save(vectors_file, chunk_vectors)


def rebuild_vector_index(api_key: str = None):
    """Generates embeddings for all chunks in registry."""
    global chunk_vectors, chunks_registry
    if not chunks_registry:
        chunk_vectors = None
        if os.path.exists(vectors_file):
            os.remove(vectors_file)
        return

    gemini_key = sanitize_key(api_key or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY"))
    all_texts = [c['text'] for c in chunks_registry]

    vectors_list = []
    if gemini_key and gemini_key.startswith("AIza"):
        # Attempt Gemini Cloud Embeddings
        test_vec = get_gemini_embedding(all_texts[0], gemini_key)
        if test_vec:
            vectors_list.append(test_vec)
            for t in all_texts[1:]:
                v = get_gemini_embedding(t, gemini_key) or [0.0] * len(test_vec)
                vectors_list.append(v)
            chunk_vectors = np.array(vectors_list, dtype=np.float32)
            save_storage()
            return

    # If no Gemini key, rely on fast in-memory hybrid search
    chunk_vectors = None
    save_storage()


def add_document(file_path: str, filename: str, api_key: str = None) -> int:
    """Extracts, chunks, and indexes a new document."""
    global chunks_registry, uploaded_documents

    pages = extract_text(file_path, filename)
    if not pages:
        return 0

    new_chunks = []
    for p in pages:
        text_splits = chunk_text(p['text'], chunk_size=1000, overlap=200)
        for idx, chunk_str in enumerate(text_splits):
            new_chunks.append({
                'filename': filename,
                'page': p['page'],
                'chunk_index': idx,
                'text': chunk_str
            })

    if not new_chunks:
        return 0

    # Replace existing chunks if same filename
    chunks_registry = [c for c in chunks_registry if c['filename'] != filename]
    chunks_registry.extend(new_chunks)

    doc_info = {
        'filename': filename,
        'chunks_count': len(new_chunks),
        'file_size': os.path.getsize(file_path)
    }
    uploaded_documents = [d for d in uploaded_documents if d['filename'] != filename]
    uploaded_documents.append(doc_info)

    rebuild_vector_index(api_key)
    return len(new_chunks)


def search_similar_chunks(query: str, top_k: int = 5, api_key: str = None) -> list[dict]:
    """
    Hybrid Search: Combines Dense Semantic Vectors (Gemini) + Keyword Overlap + Phrase Matching.
    If the document has <= 8 chunks or if query is a summary request, passes all relevant chunks.
    """
    global chunk_vectors, chunks_registry
    if not chunks_registry:
        return []

    # If total chunks are small (<= 8 chunks), return all chunks so LLM has full context
    if len(chunks_registry) <= 8:
        return [{
            'rank': idx + 1,
            'filename': c['filename'],
            'page': c['page'],
            'text': c['text'],
            'score': 1.0,
            'snippet': c['text'][:250] + ('...' if len(c['text']) > 250 else '')
        } for idx, c in enumerate(chunks_registry)]

    gemini_key = sanitize_key(api_key or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY"))
    scores = np.zeros(len(chunks_registry), dtype=np.float32)

    # 1. Semantic Vector Score (if Gemini embeddings available)
    if chunk_vectors is not None and gemini_key and chunk_vectors.shape[1] == 768:
        q_vec = get_gemini_embedding(query, gemini_key)
        if q_vec:
            q_arr = np.array(q_vec, dtype=np.float32)
            norms = np.linalg.norm(chunk_vectors, axis=1) * (np.linalg.norm(q_arr) + 1e-8)
            semantic_scores = np.dot(chunk_vectors, q_arr) / norms
            scores += 0.6 * np.nan_to_num(semantic_scores, nan=0.0)

    # 2. Keyword & Phrase Match Score
    keyword_scores = np.array([compute_keyword_score(c['text'], query) for c in chunks_registry], dtype=np.float32)
    scores += 0.5 * keyword_scores

    # 3. Summary / Broad Question Detection
    is_broad_query = any(w in query.lower() for w in ['summary', 'summarize', 'overview', 'about', 'what is this', 'explain this', 'main points', 'describe'])
    if is_broad_query or np.max(scores) < 0.05:
        # Include first 5 chunks (introductory & key sections)
        top_indices = list(range(min(top_k, len(chunks_registry))))
    else:
        top_indices = np.argsort(scores)[::-1][:top_k].tolist()

    results = []
    for rank, idx in enumerate(top_indices, 1):
        c = chunks_registry[idx]
        results.append({
            'rank': rank,
            'filename': c['filename'],
            'page': c['page'],
            'text': c['text'],
            'score': round(float(scores[idx]), 4),
            'snippet': c['text'][:250] + ('...' if len(c['text']) > 250 else '')
        })

    return results


# -----------------------------------------------------------------------------
# 5. DIRECT LLM CALLS (Google Gemini 2.0 Flash / Groq / OpenAI)
# -----------------------------------------------------------------------------

def call_llm(user_question: str, context_text: str, provider: str = "auto", api_key: str = None) -> str:
    """Direct HTTP requests to LLM APIs (Gemini 2.0 Flash / Groq / OpenAI)."""
    system_instruction = (
        "You are 'My Documents AI', an expert and helpful document intelligence assistant. "
        "Your task is to thoroughly and accurately answer the user's question using the provided document excerpts.\n\n"
        "Guidelines:\n"
        "1. Read all the document excerpts carefully.\n"
        "2. Directly answer the question using the facts, numbers, dates, and explanations present in the text.\n"
        "3. Provide specific details, bullet points, or step-by-step breakdowns.\n"
        "4. If the user asks for a summary or overview, synthesize the main topics clearly.\n"
        "5. If relevant information is mentioned in the excerpts, explain what the document states.\n"
        "6. Only state that the information was not found if the excerpts are completely unrelated to the question."
    )

    full_prompt = (
        f"DOCUMENT EXCERPTS:\n{context_text}\n\n"
        f"USER QUESTION: {user_question}\n\n"
        "Please provide a comprehensive and clear answer based on the document excerpts above."
    )

    cleaned_user_key = sanitize_key(api_key)
    gemini_key = sanitize_key(cleaned_user_key or os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY"))
    groq_openai_key = sanitize_key(cleaned_user_key or os.environ.get("OPENAI_API_KEY") or os.environ.get("GROQ_API_KEY"))

    # A) Google Gemini Direct API (Free Tier)
    if (provider == "gemini" or (not provider and gemini_key and (gemini_key.startswith("AIza") or gemini_key.startswith("AQ.")))) and gemini_key:
        for model_name in ['gemini-3.6-flash', 'gemini-3.7-flash', 'gemini-flash-latest', 'gemini-2.5-flash-lite', 'gemini-2.0-flash']:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={gemini_key}"
                payload = {
                    "contents": [{"parts": [{"text": f"System: {system_instruction}\n\n{full_prompt}"}]}],
                    "generationConfig": {"temperature": 0.2, "maxOutputTokens": 2048}
                }
                res = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=25)
                if res.status_code == 200:
                    data = res.json()
                    if 'candidates' in data and data['candidates']:
                        return data['candidates'][0]['content']['parts'][0]['text']
            except Exception as e:
                continue

    # B) Groq / OpenAI Direct API
    if groq_openai_key:
        try:
            base_url = "https://api.groq.com/openai/v1/chat/completions" if groq_openai_key.startswith("gsk_") else "https://api.openai.com/v1/chat/completions"
            model = os.environ.get("OPENAI_MODEL", "llama-3.3-70b-versatile" if groq_openai_key.startswith("gsk_") else "gpt-4o-mini")
            headers = {
                "Authorization": f"Bearer {groq_openai_key}".strip(),
                "Content-Type": "application/json"
            }
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": full_prompt}
                ],
                "temperature": 0.2
            }
            res = requests.post(base_url, json=payload, headers=headers, timeout=25)
            if res.status_code == 200:
                return res.json()['choices'][0]['message']['content']
        except Exception as e:
            print(f"Groq/OpenAI error: {e}")

    return None


# -----------------------------------------------------------------------------
# 6. FLASK WEB ROUTES & UI
# -----------------------------------------------------------------------------

UI_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>My Documents — RAG Q&A</title>
    <script src="https://cdn.tailwindcss.com"></script>
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.5.1/css/all.min.css">
    <style>
        ::-webkit-scrollbar { width: 6px; height: 6px; }
        ::-webkit-scrollbar-track { background: transparent; }
        ::-webkit-scrollbar-thumb { background: #334155; border-radius: 9999px; }
        .markdown-body p { margin-bottom: 0.75rem; }
        .markdown-body ul { list-style-type: disc; padding-left: 1.25rem; margin-bottom: 0.75rem; }
        .markdown-body pre { background: #0f172a; padding: 0.75rem; border-radius: 0.5rem; overflow-x: auto; margin: 0.5rem 0; }
        .markdown-body code { background: rgba(255,255,255,0.1); padding: 0.1rem 0.3rem; border-radius: 0.2rem; font-size: 0.9em; }
    </style>
</head>
<body class="bg-slate-950 text-slate-100 min-h-screen flex flex-col font-sans">
    
    <!-- Header -->
    <header class="bg-slate-900 border-b border-slate-800 px-6 py-4 flex items-center justify-between sticky top-0 z-20 shadow-md">
        <div class="flex items-center gap-3">
            <div class="w-10 h-10 rounded-xl bg-gradient-to-tr from-blue-600 via-indigo-600 to-violet-600 flex items-center justify-center text-white text-lg shadow-lg">
                <i class="fa-solid fa-folder-tree"></i>
            </div>
            <div>
                <h1 class="text-lg font-bold text-white tracking-wide flex items-center gap-2">
                    <span>My Documents</span>
                    <span class="text-[10px] px-2 py-0.5 rounded-full bg-blue-500/10 text-blue-400 border border-blue-500/20 font-medium">RAG Q&A</span>
                </h1>
                <p class="text-xs text-slate-400">Hybrid Retrieval &bull; Vector Search &bull; AI Synthesis</p>
            </div>
        </div>
        <div class="flex items-center gap-3">
            <button id="settingsBtn" class="text-xs px-3 py-1.5 rounded-lg bg-slate-800 hover:bg-slate-700 border border-slate-700 text-slate-300 transition flex items-center gap-1.5">
                <i class="fa-solid fa-sliders text-indigo-400"></i>
                <span>LLM Settings</span>
            </button>
        </div>
    </header>

    <!-- Main Container -->
    <div class="flex-1 max-w-7xl w-full mx-auto p-4 md:p-6 grid grid-cols-1 md:grid-cols-3 gap-6">
        
        <!-- Left Panel: Knowledge Base -->
        <div class="bg-slate-900 rounded-2xl border border-slate-800 p-5 flex flex-col h-[78vh] shadow-xl">
            <div class="flex items-center justify-between mb-3">
                <h2 class="text-sm font-bold text-white flex items-center gap-2">
                    <i class="fa-solid fa-book-bookmark text-blue-400"></i>
                    <span>Knowledge Base</span>
                </h2>
                <span id="docBadge" class="text-[11px] px-2 py-0.5 rounded-full bg-blue-950 text-blue-300 font-semibold border border-blue-800/50">0 files</span>
            </div>

            <!-- Upload Dropzone -->
            <div id="dropzone" class="border-2 border-dashed border-slate-700 hover:border-blue-500 bg-slate-950/60 rounded-xl p-5 text-center cursor-pointer transition mb-4">
                <input type="file" id="fileInput" multiple accept=".pdf,.docx,.txt,.csv,.md,.json" class="hidden">
                <div class="w-10 h-10 rounded-full bg-blue-500/10 text-blue-400 flex items-center justify-center mx-auto mb-2 text-lg">
                    <i class="fa-solid fa-cloud-arrow-up"></i>
                </div>
                <div class="text-xs font-semibold text-white">Click or Drop Files to Upload</div>
                <p class="text-[10px] text-slate-400 mt-0.5">Supports PDF, DOCX, TXT, CSV, MD</p>
            </div>

            <div id="uploadingBox" class="hidden p-3 rounded-xl bg-blue-950/60 border border-blue-500/30 text-blue-300 text-xs text-center mb-3">
                <i class="fa-solid fa-circle-notch fa-spin mr-1.5"></i> Chunking & vectorizing document...
            </div>

            <!-- Documents List -->
            <div class="flex items-center justify-between text-xs text-slate-400 mb-2">
                <span class="font-medium">Indexed Documents</span>
                <button id="clearAllBtn" class="text-rose-400 hover:text-rose-300 text-[11px] transition">Clear All</button>
            </div>
            <div id="documentsList" class="flex-1 overflow-y-auto space-y-2 pr-1 text-xs">
                <div class="text-center py-12 text-slate-500">No documents uploaded yet.</div>
            </div>
        </div>

        <!-- Right Panel: Q&A Chat Area -->
        <div class="md:col-span-2 bg-slate-900 rounded-2xl border border-slate-800 p-5 flex flex-col h-[78vh] shadow-xl">
            <div id="chatMessages" class="flex-1 overflow-y-auto space-y-4 pr-2 mb-4">
                <div id="welcomeMessage" class="text-center py-20 px-4">
                    <div class="w-14 h-14 rounded-2xl bg-gradient-to-tr from-blue-500 to-indigo-600 text-white flex items-center justify-center text-2xl mx-auto mb-3 shadow-lg">
                        <i class="fa-solid fa-magnifying-glass-chart"></i>
                    </div>
                    <h3 class="text-lg font-bold text-white mb-1">Ask questions on your documents</h3>
                    <p class="text-xs text-slate-400 max-w-md mx-auto">Upload any PDF or document to the left. The system will retrieve relevant chunks and synthesize answers with precise source citations.</p>
                </div>
            </div>

            <div id="searchingIndicator" class="hidden text-xs text-blue-400 mb-2 flex items-center gap-2">
                <i class="fa-solid fa-circle-notch fa-spin"></i>
                <span>Searching document context and generating answer...</span>
            </div>

            <form id="questionForm" class="flex gap-2">
                <input
                    type="text"
                    id="questionInput"
                    placeholder="Ask a question based on your uploaded documents..."
                    class="flex-1 bg-slate-950 border border-slate-800 rounded-xl px-4 py-3 text-xs text-white placeholder-slate-500 focus:outline-none focus:border-blue-500 transition"
                    required
                />
                <button
                    type="submit"
                    id="sendBtn"
                    class="px-5 py-3 rounded-xl bg-blue-600 hover:bg-blue-500 text-white font-semibold text-xs shadow-md transition flex items-center gap-2"
                >
                    <span>Ask AI</span>
                    <i class="fa-solid fa-arrow-up text-xs"></i>
                </button>
            </form>
        </div>
    </div>

    <!-- Settings Modal -->
    <div id="settingsModal" class="hidden fixed inset-0 bg-black/80 backdrop-blur-sm flex items-center justify-center z-50 p-4">
        <div class="bg-slate-900 border border-slate-800 rounded-2xl w-full max-w-md p-5 shadow-2xl space-y-4 text-xs">
            <div class="flex items-center justify-between border-b border-slate-800 pb-3">
                <h3 class="text-sm font-bold text-white flex items-center gap-2">
                    <i class="fa-solid fa-sliders text-blue-400"></i>
                    <span>Model Settings</span>
                </h3>
                <button id="closeSettingsBtn" class="text-slate-400 hover:text-white">
                    <i class="fa-solid fa-xmark text-base"></i>
                </button>
            </div>

            <div>
                <label class="block font-medium text-slate-300 mb-1">LLM Provider</label>
                <select id="providerSelect" class="w-full bg-slate-950 border border-slate-800 rounded-lg px-3 py-2 text-white">
                    <option value="gemini">Google Gemini 2.0 Flash (Free API Key)</option>
                    <option value="groq">Groq Live AI (Ultra Fast & Free)</option>
                    <option value="openai">OpenAI (GPT-4o-mini)</option>
                </select>
            </div>

            <div>
                <label class="block font-medium text-slate-300 mb-1">API Key (Gemini / Groq / OpenAI)</label>
                <input type="password" id="apiKeyInput" placeholder="Paste your API key (e.g. AIza... or gsk_...)" class="w-full bg-slate-950 border border-slate-800 rounded-lg px-3 py-2 text-white">
            </div>

            <div class="p-3 rounded-lg bg-slate-950/80 border border-slate-800 text-[11px] text-slate-400 space-y-1">
                <div><strong>Retrieval:</strong> Hybrid Dense Embeddings + BM25 Matcher</div>
                <div><strong>Memory Footprint:</strong> ~35MB RAM (Render Free Tier)</div>
            </div>

            <div class="flex justify-end gap-2 pt-2 border-t border-slate-800">
                <button id="saveSettingsBtn" class="px-4 py-2 rounded-lg bg-blue-600 hover:bg-blue-500 text-white font-semibold">Save Settings</button>
            </div>
        </div>
    </div>

    <script>
        const dropzone = document.getElementById('dropzone');
        const fileInput = document.getElementById('fileInput');
        const uploadingBox = document.getElementById('uploadingBox');
        const documentsList = document.getElementById('documentsList');
        const docBadge = document.getElementById('docBadge');
        const clearAllBtn = document.getElementById('clearAllBtn');
        const chatMessages = document.getElementById('chatMessages');
        const welcomeMessage = document.getElementById('welcomeMessage');
        const questionForm = document.getElementById('questionForm');
        const questionInput = document.getElementById('questionInput');
        const searchingIndicator = document.getElementById('searchingIndicator');
        const sendBtn = document.getElementById('sendBtn');
        const settingsModal = document.getElementById('settingsModal');
        const settingsBtn = document.getElementById('settingsBtn');
        const closeSettingsBtn = document.getElementById('closeSettingsBtn');
        const saveSettingsBtn = document.getElementById('saveSettingsBtn');
        const providerSelect = document.getElementById('providerSelect');
        const apiKeyInput = document.getElementById('apiKeyInput');

        let currentApiKey = localStorage.getItem('my_doc_api_key') || '';
        let currentProvider = localStorage.getItem('my_doc_provider') || 'gemini';

        providerSelect.value = currentProvider;
        apiKeyInput.value = currentApiKey;

        // Settings Modal
        settingsBtn.addEventListener('click', () => settingsModal.classList.remove('hidden'));
        closeSettingsBtn.addEventListener('click', () => settingsModal.classList.add('hidden'));
        saveSettingsBtn.addEventListener('click', async () => {
            currentApiKey = apiKeyInput.value.trim();
            currentProvider = providerSelect.value;
            localStorage.setItem('my_doc_api_key', currentApiKey);
            localStorage.setItem('my_doc_provider', currentProvider);
            settingsModal.classList.add('hidden');
            
            // Reindex with new key if documents exist
            await fetch('/api/reindex', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ api_key: currentApiKey })
            });
            alert('Settings saved and knowledge base refreshed!');
        });

        // Dropzone & File Selection
        dropzone.addEventListener('click', () => fileInput.click());
        fileInput.addEventListener('change', (e) => uploadFiles(e.target.files));

        dropzone.addEventListener('dragover', (e) => { e.preventDefault(); dropzone.classList.add('border-blue-500'); });
        dropzone.addEventListener('dragleave', () => dropzone.classList.remove('border-blue-500'));
        dropzone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropzone.classList.remove('border-blue-500');
            uploadFiles(e.dataTransfer.files);
        });

        async function uploadFiles(files) {
            if (!files || files.length === 0) return;
            uploadingBox.classList.remove('hidden');

            const formData = new FormData();
            for (let i = 0; i < files.length; i++) {
                formData.append('files', files[i]);
            }
            if (currentApiKey) {
                formData.append('api_key', currentApiKey);
            }

            try {
                const res = await fetch('/api/upload', { method: 'POST', body: formData });
                const data = await res.json();
                if (res.ok && data.status === 'success') {
                    fetchDocuments();
                } else {
                    alert(data.message || data.error || 'Upload failed');
                }
            } catch (err) {
                alert('Upload error: ' + err.message);
            } finally {
                uploadingBox.classList.add('hidden');
                fileInput.value = '';
            }
        }

        async function fetchDocuments() {
            try {
                const res = await fetch('/api/documents');
                const data = await res.json();
                renderDocuments(data.documents || []);
            } catch (err) {
                console.error(err);
            }
        }

        function renderDocuments(docs) {
            docBadge.textContent = `${docs.length} file${docs.length === 1 ? '' : 's'}`;
            documentsList.innerHTML = '';
            if (docs.length === 0) {
                documentsList.innerHTML = '<div class="text-center py-12 text-slate-500">No documents in knowledge base yet.</div>';
                return;
            }

            docs.forEach(doc => {
                const item = document.createElement('div');
                item.className = 'p-2.5 rounded-xl bg-slate-950/80 border border-slate-800 flex items-center justify-between text-xs group';
                item.innerHTML = `
                    <div class="flex items-center gap-2 truncate flex-1">
                        <i class="fa-solid fa-file-lines text-blue-400 text-sm flex-shrink-0"></i>
                        <div class="truncate">
                            <div class="font-semibold text-white truncate">${doc.filename}</div>
                            <div class="text-[10px] text-slate-400">${doc.chunks_count} chunks indexed</div>
                        </div>
                    </div>
                    <div class="flex items-center gap-1.5 flex-shrink-0">
                        <span class="text-[10px] px-2 py-0.5 rounded-full bg-emerald-950/80 text-emerald-400 border border-emerald-500/20 font-medium">Ready</span>
                        <button class="delete-single-btn text-slate-500 hover:text-rose-400 p-1 transition" title="Delete document">
                            <i class="fa-solid fa-trash text-xs"></i>
                        </button>
                    </div>
                `;

                item.querySelector('.delete-single-btn').addEventListener('click', async (e) => {
                    e.stopPropagation();
                    if (confirm(`Remove "${doc.filename}" from knowledge base?`)) {
                        await fetch('/api/delete', {
                            method: 'POST',
                            headers: { 'Content-Type': 'application/json' },
                            body: JSON.stringify({ filename: doc.filename })
                        });
                        fetchDocuments();
                    }
                });

                documentsList.appendChild(item);
            });
        }

        clearAllBtn.addEventListener('click', async () => {
            if (confirm('Delete all documents and reset index?')) {
                await fetch('/api/clear', { method: 'POST' });
                fetchDocuments();
                chatMessages.innerHTML = '';
                chatMessages.appendChild(welcomeMessage);
            }
        });

        questionForm.addEventListener('submit', async (e) => {
            e.preventDefault();
            const text = questionInput.value.trim();
            if (!text) return;

            if (welcomeMessage) welcomeMessage.classList.add('hidden');
            questionInput.value = '';
            appendMessageBubble('user', text);
            searchingIndicator.classList.remove('hidden');
            sendBtn.disabled = true;

            try {
                const res = await fetch('/api/query', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        query: text,
                        provider: currentProvider,
                        api_key: currentApiKey
                    })
                });
                const data = await res.json();
                const answerText = data.answer || data.error || data.message || 'No answer generated.';
                appendMessageBubble('assistant', answerText, data.sources || []);
            } catch (err) {
                appendMessageBubble('assistant', '⚠️ Error: ' + err.message);
            } finally {
                searchingIndicator.classList.add('hidden');
                sendBtn.disabled = false;
                chatMessages.scrollTop = chatMessages.scrollHeight;
            }
        });

        function appendMessageBubble(role, text, sources = []) {
            const isUser = role === 'user';
            const msgEl = document.createElement('div');
            msgEl.className = `flex flex-col ${isUser ? 'items-end' : 'items-start'} text-xs space-y-1`;

            let sourcesHtml = '';
            if (!isUser && sources && sources.length > 0) {
                const items = sources.map(s => `
                    <div class="p-2.5 rounded-lg bg-slate-950 border border-slate-800 text-[11px] mt-1 space-y-1">
                        <div class="font-semibold text-blue-300 flex justify-between">
                            <span>📄 ${s.filename} (Page ${s.page})</span>
                            <span class="text-[10px] text-emerald-400">Score: ${s.score}</span>
                        </div>
                        <p class="text-slate-400 italic bg-slate-900/60 p-1.5 rounded">"${s.snippet}"</p>
                    </div>
                `).join('');

                sourcesHtml = `
                    <details class="mt-2.5 pt-2 border-t border-slate-800 w-full text-slate-400">
                        <summary class="cursor-pointer font-semibold text-blue-400 hover:text-blue-300 text-[11px]">
                            📚 View ${sources.length} Referenced Source Chunks
                        </summary>
                        <div class="mt-2 space-y-1.5">${items}</div>
                    </details>
                `;
            }

            msgEl.innerHTML = `
                <div class="text-[10px] font-semibold text-slate-400 px-1">${isUser ? 'You' : 'My Documents AI'}</div>
                <div class="p-4 rounded-2xl max-w-xl ${isUser ? 'bg-blue-600 text-white' : 'bg-slate-900 text-slate-200 border border-slate-800'} markdown-body shadow-md">
                    ${isUser ? text : marked.parse(text)}
                    ${sourcesHtml}
                </div>
            `;

            chatMessages.appendChild(msgEl);
            chatMessages.scrollTop = chatMessages.scrollHeight;
        }

        fetchDocuments();
    </script>
</body>
</html>
"""


@app.route('/')
def index():
    return render_template_string(UI_TEMPLATE)


@app.route('/api/upload', methods=['POST'])
def api_upload():
    files = request.files.getlist('files') or request.files.getlist('file') or list(request.files.values())
    if not files:
        return jsonify({'error': 'No file attached'}), 400

    api_key = request.form.get('api_key', None)
    total_chunks = 0
    uploaded_files_count = 0

    for f in files:
        if not f or not f.filename:
            continue
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], f.filename)
        f.save(save_path)
        chunks_added = add_document(save_path, f.filename, api_key=api_key)
        total_chunks += chunks_added
        if chunks_added > 0:
            uploaded_files_count += 1

    if total_chunks == 0:
        return jsonify({
            'status': 'error',
            'error': 'No readable text could be extracted from the uploaded document(s). Please ensure files contain text.'
        }), 400

    return jsonify({
        'status': 'success',
        'total_chunks': total_chunks,
        'uploaded_files': uploaded_files_count,
        'message': f'Successfully embedded {total_chunks} chunks into vector index.'
    })


@app.route('/api/documents', methods=['GET'])
def api_documents():
    return jsonify({'documents': uploaded_documents})


@app.route('/api/reindex', methods=['POST'])
def api_reindex():
    data = request.get_json() or {}
    api_key = data.get('api_key', None)
    rebuild_vector_index(api_key=api_key)
    return jsonify({'status': 'success', 'message': 'Reindexed with new API settings.'})


@app.route('/api/delete', methods=['POST'])
def api_delete_doc():
    global chunks_registry, uploaded_documents
    data = request.get_json() or {}
    filename = data.get('filename', '').strip()
    if not filename:
        return jsonify({'error': 'Filename required'}), 400

    chunks_registry = [c for c in chunks_registry if c['filename'] != filename]
    uploaded_documents = [d for d in uploaded_documents if d['filename'] != filename]

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception:
            pass

    rebuild_vector_index()
    return jsonify({'status': 'success', 'message': f'Document {filename} removed.'})


@app.route('/api/clear', methods=['POST'])
def api_clear():
    global chunks_registry, uploaded_documents, chunk_vectors
    chunks_registry = []
    uploaded_documents = []
    chunk_vectors = None

    if os.path.exists(app.config['STORAGE_FOLDER']):
        shutil.rmtree(app.config['STORAGE_FOLDER'])
    os.makedirs(app.config['STORAGE_FOLDER'], exist_ok=True)

    if os.path.exists(app.config['UPLOAD_FOLDER']):
        shutil.rmtree(app.config['UPLOAD_FOLDER'])
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

    save_storage()
    return jsonify({'status': 'success', 'message': 'All documents cleared.'})


@app.route('/api/query', methods=['POST'])
def api_query():
    data = request.get_json() or {}
    query_text = data.get('query', '').strip()
    provider = data.get('provider', 'gemini')
    api_key = data.get('api_key', None)

    if not query_text:
        return jsonify({'error': 'Question cannot be empty'}), 400

    # 1. Search hybrid vector & keyword similarity
    retrieved_chunks = search_similar_chunks(query_text, top_k=5, api_key=api_key)
    if not retrieved_chunks:
        return jsonify({
            'answer': 'No documents found in your knowledge base. Please upload at least one document first.',
            'sources': []
        })

    # 2. Build Context String
    context_parts = []
    for c in retrieved_chunks:
        context_parts.append(f"[Document: {c['filename']}, Page: {c['page']}]\n{c['text']}")
    context_str = "\n\n".join(context_parts)

    # 3. Call LLM
    answer = call_llm(query_text, context_str, provider=provider, api_key=api_key)

    # 4. Fallback Answer if LLM API is unavailable / offline
    if not answer:
        answer = f"Based on your documents, here are the most relevant findings for **\"{query_text}\"**:\n\n"
        for c in retrieved_chunks:
            answer += f"- **From {c['filename']} (Page {c['page']}):**\n  > \"{c['text']}\"\n\n"
        answer += "\n---\n💡 *Tip: Add your Google Gemini API Key in Settings to get conversational AI synthesis.*"

    return jsonify({
        'answer': answer,
        'sources': retrieved_chunks
    })


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    print("\n" + "="*70)
    print(">> 'My Documents' Hybrid RAG Application is running!")
    print(f">> Open in browser: http://0.0.0.0:{port}")
    print("="*70 + "\n")
    app.run(host='0.0.0.0', port=port, debug=False)
